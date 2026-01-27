import os
from datetime import datetime
from typing import List
from fastapi import (
    UploadFile, File, Form, Depends, HTTPException, APIRouter, status, Request)
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import func
from docx import Document
from docx.shared import Inches
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.config import GENERATED_FOLDER
from app.db.database import get_db
from app.models.rfp_models import User, RFPDocument,GeneratedRFPDocument,RFPQuestion
from app.api.routes.utils import get_current_user
from app.utils.admin_function import upload_documents,get_final_answer,clean_text,add_footer_page_numbers,add_formatted_text,extract_question_number

router = APIRouter()

@router.post("/generate-rfp-doc/")
async def generate_rfp_doc(
    rfp_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role.lower() != "admin":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Unauthorized"
        )

    rfp_doc = db.query(RFPDocument).filter(
        RFPDocument.id == rfp_id,
        RFPDocument.is_deleted == False
    ).first()

    if not rfp_doc:
        raise HTTPException(status_code=404, detail="RFP not found")

    unanswered = []

    for q in rfp_doc.questions:
        if getattr(q, "is_deleted", False):
            continue

        has_reviewer_ans = any(
            rev.submit_status == "submitted"
            for rev in q.reviewers
        )
        has_ai_ans = bool(q.answer_versions)

        if not (has_reviewer_ans or has_ai_ans):
            unanswered.append(q.question_text)

    if unanswered:
        raise HTTPException(
            status_code=409,
            detail="Some questions are not answered yet",
            headers={"X-Unanswered-Count": str(len(unanswered))}
        )

    if not rfp_doc.summary:
        raise HTTPException(
            status_code=404,
            detail="Executive summary not found"
        )

    questions = (
        db.query(RFPQuestion)
        .filter(RFPQuestion.rfp_id == rfp_id)
        .order_by(
            RFPQuestion.assigned_at.asc(),
            RFPQuestion.id.asc()
        )
        .all()
    )

    last_version = db.query(func.max(GeneratedRFPDocument.version)) \
        .filter(GeneratedRFPDocument.rfp_id == rfp_id) \
        .scalar() or 0

    new_version = last_version + 1

    doc = Document()

    try:
        logo_path = "image.png"
        doc.add_picture(logo_path, width=Inches(1.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception:
        pass

    doc.add_heading("Proposal", level=0)
    doc.add_paragraph("Presented by Ringer")
    doc.add_paragraph(f"Client: {rfp_doc.project_name or 'Client'}")
    doc.add_paragraph(
        f"Generated on {datetime.utcnow().strftime('%d %B %Y')}"
    )

    p = doc.add_paragraph("")
    p.paragraph_format.space_after = Pt(24)

    h = doc.add_heading("Executive Summary", level=1)
    h.paragraph_format.space_after = Pt(14)
    add_formatted_text(doc, rfp_doc.summary.summary_text)

    p = doc.add_paragraph("")
    p.paragraph_format.space_after = Pt(18)

    h = doc.add_heading("Proposal Response", level=1)
    h.paragraph_format.space_after = Pt(14)

    question_counter = 1

    for q in questions:
        if getattr(q, "is_deleted", False):
            continue

        final_answer = get_final_answer(q)
        if not final_answer:
            continue

        q_number, q_text = extract_question_number(q.question_text)

        q_para = doc.add_paragraph()
        if q_number:
            q_run = q_para.add_run(f"{q_number} {q_text}")
        else:
            q_run = q_para.add_run(f"{question_counter}.0 {q_text}")
        q_run.bold = True
        q_para.paragraph_format.space_after = Pt(12)

        add_formatted_text(doc, final_answer)

        question_counter += 1

    add_footer_page_numbers(doc)

    os.makedirs(GENERATED_FOLDER, exist_ok=True)

    base_name = os.path.splitext(rfp_doc.filename)[0]
    file_name = f"{base_name}_proposal_v{new_version}.docx"
    file_path = os.path.join(GENERATED_FOLDER, file_name)

    doc.save(file_path)

    gen_doc = GeneratedRFPDocument(
        rfp_id=rfp_id,
        file_name=file_name,
        file_path=file_path,
        generated_by=current_user.id,
        version=new_version
    )

    db.add(gen_doc)
    db.commit()
    db.refresh(gen_doc)

    base_url = str(request.base_url).rstrip("/")
    download_url = f"{base_url}/documents/{gen_doc.id}/download"

    return {
        "message": "Proposal generated successfully",
        "document_id": gen_doc.id,
        "version": new_version,
        "download_url": download_url
    }

@router.get("/list-rfp-docs/")
async def list_rfp_docs(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role.lower() != "admin":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Unauthorized"
        )

    docs = (
        db.query(GeneratedRFPDocument)
        .filter(GeneratedRFPDocument.is_deleted == False)
        .order_by(GeneratedRFPDocument.generated_at.desc())
        .all()
    )

    base_url = str(request.base_url).rstrip("/")

    return {
        "count": len(docs),
        "docs": [
            {
                "document_id": d.id,
                "rfp_id": d.rfp_id,
                "file_name": d.file_name,
                "version": d.version,
                "generated_at": d.generated_at,
                "download_url": f"{base_url}/documents/{d.id}/download"
            }
            for d in docs
        ]
    }

@router.post("/upload-library")
def upload_library_new(
    files: List[UploadFile] = File(...),
    project_name: str = Form(...),
    category: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role.lower() != "admin":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Only admins can upload library documents."
        )

    try:
        uploaded_docs = upload_documents(files, project_name, category, current_user, db)
        return {
            "message": f"{len(uploaded_docs)} file(s) uploaded successfully",
            "documents": uploaded_docs
        }

    except Exception as e:
        import traceback
        print(traceback.format_exc())
        raise HTTPException(
            status_code=409,
            detail="This RFP already exist."
        )

@router.get("/documents/{doc_id}/download")
def download_generated_document(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role.lower() != "admin":
        raise HTTPException(status_code=401, detail="Unauthorized")

    doc = db.query(GeneratedRFPDocument).filter(
        GeneratedRFPDocument.id == doc_id,
        GeneratedRFPDocument.is_deleted == False
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if not os.path.exists(doc.file_path):
        raise HTTPException(
            status_code=404,
            detail="File not found on server"
        )

    return FileResponse(
        path=doc.file_path,
        filename=doc.file_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@router.delete("/delete-gen-doc/{doc_id}")
def delete_generated_document(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role.lower() != "admin":
        raise HTTPException(status_code=401, detail="Unauthorized")

    doc = db.query(GeneratedRFPDocument).filter(
        GeneratedRFPDocument.id == doc_id,
        GeneratedRFPDocument.is_deleted == False
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)

    doc.is_deleted = True
    doc.deleted_at = datetime.utcnow()
    db.commit()

    return {
        "message": "Document deleted successfully",
        "document_id": doc_id
    }
