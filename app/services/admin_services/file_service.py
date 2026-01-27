import os
import shutil
import hashlib
from datetime import datetime
from fastapi import HTTPException
from sqlalchemy.orm import Session
from app.config import UPLOAD_FOLDER, index
from app.models.rfp_models import RFPDocument,RFPQuestion
from app.services.llm_services.llm_service import (
    extract_text_from_file,
    generate_summary,
    get_embedding
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx import Document
from docx.shared import Inches, Pt



def upload_documents(files, project_name, category, current_user, db: Session):
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    uploaded_docs = []

    for file in files:
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in [".pdf", ".docx", ".pptx", ".xlsx", ".xls"]:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file format: {file.filename}"
            )

        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
        saved_filename = f"{timestamp}_{file.filename}"
        file_path = os.path.join(UPLOAD_FOLDER, saved_filename)

        with open(file_path, "wb") as f:
            shutil.copyfileobj(file.file, f)

        with open(file_path, "rb") as f:
            file_bytes = f.read()
            file_hash = hashlib.sha256(file_bytes).hexdigest()

        new_doc = RFPDocument(
            filename=file.filename,
            file_path=file_path,
            category=category,
            project_name=project_name,
            admin_id=current_user.id,
            uploaded_at=datetime.utcnow(),
            file_hash=file_hash
        )
        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)

        text = extract_text_from_file(file_path)
        if not text:
            continue

        summary = generate_summary(text)
        summary_vector = get_embedding(summary)
        index.upsert(
            vectors=[(
                f"summary_{new_doc.id}",
                summary_vector,
                {
                    "document_id": str(new_doc.id),
                    "filename": new_doc.filename,
                    "category": new_doc.category,
                    "project_name": new_doc.project_name,
                    "type": "summary",
                    "text": summary
                }
            )],
            namespace="summaries"
        )

        splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = splitter.split_text(text)
        vectors = []
        for i, chunk in enumerate(chunks):
            vector = get_embedding(chunk)
            vectors.append((
                f"{new_doc.id}_{i}",
                vector,
                {
                    "document_id": str(new_doc.id),
                    "filename": new_doc.filename,
                    "category": new_doc.category,
                    "project_name": new_doc.project_name,
                    "type": "chunk",
                    "chunk_id": i,
                    "text": chunk
                }
            ))

        if vectors:
            index.upsert(vectors=vectors, namespace=f"rfp_{new_doc.id}")

        uploaded_docs.append({
            "document_id": new_doc.id,
            "filename": new_doc.filename,
            "category": new_doc.category,
            "project_name": new_doc.project_name
        })

    return uploaded_docs

def get_final_answer(question: RFPQuestion) -> str | None:
    for rev in question.reviewers:
        if rev.submit_status == "submitted" and rev.ans:
            return clean_text(rev.ans)

    if question.answer_versions:
        latest_ai = max(
            question.answer_versions,
            key=lambda x: x.generated_at
        )
        return clean_text(latest_ai.answer)

    return None

def clean_text(text: str) -> str:
    if not text:
        return ""
    return (
        text.replace('"', '')
            .replace('*', '')
            .strip()
    )

def bullet_line(line: str) -> bool:
    return line.strip().startswith(("-", "*", "•"))

def add_footer_page_numbers(doc: Document):
    section = doc.sections[0]
    footer = section.footer

    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para.add_run("Page ")

    run = para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(ns.qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(ns.qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    para.add_run(" of ")

    run = para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(ns.qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.text = " NUMPAGES "

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(ns.qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_formatted_text(doc: Document, text: str):
    if not text:
        return

    lines = text.split("\n")
    in_bullet_block = False

    for line in lines:
        line = clean_text(line)
        if not line:
            continue

        # -------- BULLET ----------
        if bullet_line(line):
            in_bullet_block = True
            bullet = doc.add_paragraph(
                line.lstrip("-*• ").strip(),
                style="List Bullet"
            )
            bullet.paragraph_format.left_indent = Inches(0.5)
            bullet.paragraph_format.space_after = Pt(2)

        # -------- NORMAL PARAGRAPH ----------
        else:
            in_bullet_block = False
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(10)

    end = doc.add_paragraph("")
    end.paragraph_format.space_after = Pt(8)

import re

def extract_question_number(text: str):
    """
    Extracts leading numbering like '1.1', '2.3.4' from question text.
    Returns (number, cleaned_text)
    """
    text = clean_text(text)

    match = re.match(r'^(\d+(\.\d+)+)\s+(.*)', text)
    if match:
        return match.group(1), match.group(3)

    return None, text
