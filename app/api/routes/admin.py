import os
from datetime import datetime
from typing import List
from fastapi import (
    UploadFile, File, Form, Depends, HTTPException, APIRouter, status)
from fastapi.staticfiles import StaticFiles
from fastapi_mail import FastMail, MessageSchema, MessageType
from sqlalchemy.orm import Session
from docx import Document
from collections import defaultdict
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.config import mail_config,LOGIN_URL,GENERATED_FOLDER 
from app.db.database import get_db, Base, engine
from app.schemas.schema import (
    FileDetails, AssignReviewer, ReviewerOut, AdminEditRequest,
    RFPDocumentGroupedQuestionsOut, NotificationRequest,
    reviwerdelete, ChatInputRequest, ReassignReviewerRequest, GroupedRFPQuestionOut,QuestionOut,QuestionInput,KeystoneCreateRequest,KeystoneUpdateRequest,KeystoneDynamicFormRequest,KeystonePatchRequest)
from app.models.rfp_models import User, Reviewer, RFPDocument, RFPQuestion
from app.services.llm_service import client 
from app.api.routes.utils import get_current_user
from app.utils.admin_function import (
    process_rfp_file, fetch_file_details,
    get_all_users, get_assigned_users,
    assign_multiple_review, get_reviewers_by_file_service,
    get_user_by_id_service, check_submissions_service,
    get_assign_user_status_service, delete_rfp_document_service,
    remove_user_service, filter_question_service,
    admin_filter_questions_by_status_service, analyze_overall_score_service,
    view_rfp_document_service, edit_question_by_admin_service,
    update_profile_service, delete_reviewer_service,
    regenerate_answer_with_chat_service, reassign_reviewer_service,upload_documents,add_ques,restore_rfp_doc,permanent_delete_rfp,get_trash_documents,extract_col,save_form,fetch_form,update_form,delete_form
)
from fastapi import Request
from fastapi.responses import JSONResponse

router = APIRouter()
Base.metadata.create_all(engine)


UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

router.mount("/uploads", StaticFiles(directory=UPLOAD_FOLDER), name="uploads")

@router.post("/search-related-summary/")
async def search_related_summary(
    file: UploadFile = File(...),
    project_name: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "admin":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Only admins can access summary docs."
        )
    
    return await process_rfp_file(file, project_name, db, current_user)

@router.get("/filedetails", response_model=List[FileDetails])
def get_file_details(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "admin":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Only admins can access File details."
        )
    
    return fetch_file_details(db)

@router.get("/userdetails")
def get_user(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return get_all_users(db, current_user)

@router.get("/get_assign_users")
def get_assigned_users(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return get_assigned_users(db, current_user)

@router.get("/userdetails/{user_id}")
def get_user_by_id_route(
    user_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return get_user_by_id_service(user_id, db)

@router.get("/rfpdetails/{document_id}/{status}", response_model=RFPDocumentGroupedQuestionsOut)
def get_rfp_details(
    document_id: int,
    status: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "admin":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Only admins can access RFP details."
        )

    valid_statuses = ["assigned", "unassigned", "total question"]
    if status not in valid_statuses:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid status. Allowed values: {valid_statuses}"
        )

    document = (
        db.query(RFPDocument)
        .filter(RFPDocument.id == document_id)
        .first()
    )

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="RFP Document not found or access denied."
        )

    grouped = defaultdict(list)

    for q in sorted(document.questions, key=lambda x: x.id):
        reviewers = db.query(Reviewer).filter(Reviewer.ques_id == q.id).all()

        if status == "assigned" and not reviewers:
            continue
        elif status == "unassigned" and reviewers:
            continue
        elif status == "total-question":
            pass  # include all questions

        grouped[q.section].append({
            "id": q.id,
            "question_text": q.question_text
        })

    grouped_questions = [
        GroupedRFPQuestionOut(
            section=section,
            questions=[QuestionOut(**q) for q in questions]
        )
        for section, questions in grouped.items()
    ]

    return {
        "id": document.id,
        "filename": document.filename,
        "uploaded_at": document.uploaded_at,
        "summary": document.summary,
        "questions_by_section": grouped_questions
    }

@router.post("/assign-reviewer")
def assign_multiple_reviewers(
    request: AssignReviewer,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return assign_multiple_review(request, db, current_user)

@router.get("/assigned-reviewers/{file_id}", response_model=list[ReviewerOut])
def get_reviewers_by_file(file_id: int, db: Session = Depends(get_db)):
    return get_reviewers_by_file_service(file_id, db)

@router.post("/send-assignment-notification")
async def send_assignment_notification_bulk(
    request: NotificationRequest, db: Session = Depends(get_db)
):
    try:
        fm = FastMail(mail_config)
        summary = []

        for uid in request.user_id:
            user = db.query(User).filter(User.id == uid).first()
            if not user or not user.email:
                continue

            questions = []
            for ques_id in request.ques_ids:
                assignment = db.query(Reviewer).filter(
                    Reviewer.ques_id == ques_id,
                    Reviewer.user_id == uid
                ).first()
                question = db.query(RFPQuestion).filter(RFPQuestion.id == ques_id).first()

                if assignment and question:
                    questions.append((question, assignment))

            if not questions:
                continue

            question_texts = "<br><br>".join([
                f"<b>QID:</b> {q.id}<br><b>Section:</b> {q.section or 'N/A'}<br><b>Question:</b> {q.question_text}"
                for q, _ in questions
            ])

            html_body = f"""
                <p>Hello {user.username},</p>

                <p>The following questions have been assigned to you:</p>

                {question_texts}

                <p>Please click the button below to log in and review:</p>

                <a href="{LOGIN_URL}" 
                   style="display:inline-block; padding:10px 20px; font-size:16px; 
                          color:#fff; background-color:#007BFF; text-decoration:none; 
                          border-radius:5px;">
                    Log In
                </a>

                <p>Best regards,<br>RFP Automation System</p>
            """

            message = MessageSchema(
                subject="Multiple RFP Questions Assigned",
                recipients=[user.email],
                body=html_body,
                subtype=MessageType.html
            )

            await fm.send_message(message)

            for _, assignment in questions:
                assignment.status = "notified"

            summary.append({
                "user_id": uid,
                "email": user.email,
                "notified_questions": [q.id for q, _ in questions]
            })

        db.commit()

        return {
            "message": "Notification emails sent successfully",
            "notifications": summary
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to send emails: {str(e)}")

@router.get("/check_submit")
def check(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return check_submissions_service(db, current_user)

@router.get("/assign_user_status")
def assign_status(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return get_assign_user_status_service(db, current_user)
  
@router.delete("/rfp/{rfp_id}")
def delete_rfp_document(
    rfp_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return delete_rfp_document_service(rfp_id, db, current_user)

@router.delete("/reviewer-remove")
async def remove_user(
    ques_id: int,
    user_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return await remove_user_service(ques_id, user_id, db, current_user)

@router.get("/filter/{rfp_id}")
def filter_question(
    rfp_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return filter_question_service(rfp_id, db, current_user)

@router.get("/admin/filter-questions-by-user/{status}")
def admin_filter_questions_by_status(
    status: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return admin_filter_questions_by_status_service(status, db, current_user)
 
@router.post("/admin/analyze-answers")
def analyze_overall_score_only_if_complete(
    rfp_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role.lower() != "admin":
        raise HTTPException(status_code=403, detail="Only admins can access this endpoint.")
    
    return analyze_overall_score_service(rfp_id, db)

@router.get("/rfp-documents/{rfp_id}/view")
def view_rfp_document(
    rfp_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "admin":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Only admins can view documents."
        )

    return view_rfp_document_service(rfp_id, db)

@router.post("/generate-rfp-doc/")
async def generate_rfp_doc(
    rfp_id: int,
    request: Request, 
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role.lower() != "admin":
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Unauthorized")
    
    rfp_doc = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    if not rfp_doc:
        raise HTTPException(status_code=404, detail="RFP not found")

    unanswered = []
    for q in rfp_doc.questions:
        has_reviewer_ans = any(rev.submit_status == "submitted" for rev in q.reviewers)
        has_ai_ans = bool(q.answer_versions)
        if not (has_reviewer_ans or has_ai_ans):
            unanswered.append(q.question_text)

    if unanswered:
        raise HTTPException(
            status_code=409,
            detail={
                "message": f"Report cannot be generated. {len(unanswered)} question(s) are not analyzed yet.",
                "unanswered_questions": unanswered
            }
        )

    summary_obj = rfp_doc.summary
    if not summary_obj:
        raise HTTPException(status_code=404, detail="Executive summary not found")
    executive_summary = summary_obj.summary_text

    company_name = getattr(rfp_doc, "client_name", "Ringer")

    doc = Document()

    try:
        logo_path = "image.png" 
        doc.add_picture(logo_path, width=Inches(1.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception as e:
        print("Logo could not be added:", e)

    doc.add_heading("RFP Proposal Response", level=0)
    doc.add_paragraph(f"Presented by Ringer")
    doc.add_paragraph(f"Client: {company_name}")
    doc.add_paragraph(f"Generated on {datetime.utcnow().strftime('%Y-%m-%d')}")

    doc.add_page_break()
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(executive_summary)

    os.makedirs(GENERATED_FOLDER, exist_ok=True)
    file_name = f"rfp_response_{rfp_id}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.docx"
    file_path = os.path.join(GENERATED_FOLDER, file_name)
    doc.save(file_path)

    base_url = str(request.base_url).rstrip("/")
    download_url = f"{base_url}/download/{file_name}"

    return {
        "message": "RFP proposal generated successfully",
        "download_url": download_url
    }

@router.get("/list-rfp-docs/")
async def list_rfp_docs(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        if current_user.role.lower() != "admin":
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Unauthorized - only admins can access this endpoint"
            )

        if not os.path.exists(GENERATED_FOLDER):
            return JSONResponse(content={
                "message": "No documents found",
                "role": current_user.role,
                "docs": []
            })

        files = os.listdir(GENERATED_FOLDER)
        files = [f for f in files if f.endswith(".docx") or f.endswith(".pdf")]

        docs = []
        for f in files:
            docs.append({
                "file_name": f,
                "download_url": f"{request.base_url}download/{f}"
            })

        return {
            "message": f"{len(docs)} document(s) found",
            "role": current_user.role,
            "docs": docs
        }

    except HTTPException as e:
        raise e
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "message": "An error occurred while listing documents",
                "error": str(e),
                "docs": []
            }
        )

@router.patch("/admin/edit-answer")
def edit_question_by_admin(
    request: AdminEditRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "admin":
        raise HTTPException(
            status_code=403,
            detail="Only admins can edit submitted questions."
        )

    try:
        return edit_question_by_admin_service(request, db)
    except HTTPException as http_exc:
        raise http_exc
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.put("/update-profile")
async def update_profile(
    username: str = Form(...),
    email: str = Form(...),
    image: UploadFile = File(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    try:
        return await update_profile_service(db, current_user, username, email, image)
    except HTTPException as http_exc:
        raise http_exc
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/delete-reviewer_user")
async def delete_reviewer(
    request: reviwerdelete,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "admin":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Only admins can remove reviewers."
        )

    if request.role == "admin":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Admin cannot be deleted"
        )

    return await delete_reviewer_service(request, db)

@router.post("/questions/chat_input")
async def regenerate_answer_with_chat(
    request: ChatInputRequest,
    db: Session = Depends(get_db),
):
    try:
        return await regenerate_answer_with_chat_service(request, db)
    except HTTPException as http_exc:
        raise http_exc
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/reassign")
async def reassign_reviewer(
    request: ReassignReviewerRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "admin":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Only admins can reassign reviewers."
        )

    try:
        return await reassign_reviewer_service(request, db, current_user)
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/upload-library")
def upload_library_new(
    files: List[UploadFile] = File(...),
    project_name: str = Form(...),
    category: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role.lower() != "admin":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Only admins can upload library documents."
        )

    try:
        uploaded_docs = upload_documents(files, project_name, category, current_user, db)
        return {
            "message": f"{len(uploaded_docs)} file(s) uploaded successfully",
            "documents": uploaded_docs
        }

    except Exception as e:
        import traceback
        print(traceback.format_exc())
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e)
        )
    
@router.post("/add/questions/{rfp_id}")
def add_questions(
    rfp_id: int,
    request: QuestionInput,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return add_ques(rfp_id, request, db, current_user)

@router.post("/rfp/{rfp_id}/restore")
def restore_rfp_document(
    rfp_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return restore_rfp_doc(rfp_id, db, current_user)

@router.delete("/rfp/{rfp_id}/permanent")
def permanent_delete_doc(
    rfp_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return permanent_delete_rfp(rfp_id, db, current_user)

@router.get("/rfp/trash")
def get_trash_doc(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return get_trash_documents(db, current_user)

# @router.post("/createform")
# def create_keystone_field(
#     request: KeystoneCreateRequest,
#     db: Session = Depends(get_db),
#     current_user: User = Depends(get_current_user)
# ):
#     return create_keystone(request, db,current_user)

# @router.get("/form")
# def fetch_keystone_form(
#     db: Session = Depends(get_db),
#     current_user: User = Depends(get_current_user)
# ):
    
#     return get_keystone_form(db,current_user)

# @router.put("/form/update/{field_id}")
# def update_form_field(
#     field_id: int,
#     request: KeystoneCreateRequest,
#     db: Session = Depends(get_db),
#     current_user: User = Depends(get_current_user)
# ):

#     return modify_form_field(field_id, request, db,current_user)

# @router.patch("/form/update/{field_id}")
# def partial_update(
#     field_id: int,
#     request: KeystoneUpdateRequest,
#     db: Session = Depends(get_db),
#     current_user: User = Depends(get_current_user)
# ):
#     return modify_partial_form(field_id, request, db,current_user)

# @router.delete("/form/delete/{field_id}")
# def delete_form(
#     field_id: int,
#     db: Session = Depends(get_db),
#     current_user: User = Depends(get_current_user)
# ):
#     return delete_keystone(field_id, db,current_user)


#spreadheet upload implementation
@router.post("/dynamic/extract-columns")
async def extract_columns(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    try:
        return await extract_col(file,current_user)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/dynamic/save-form")
def save_keystone_form(
    request: KeystoneDynamicFormRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return save_form(request, db, current_user)

@router.get("/dynamic/get-form")
def get_keystone_form(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return fetch_form(db, current_user)

@router.patch("/dynamic/form/{form_id}")
def update_keystone_form(
    form_id: int,
    request: KeystonePatchRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return update_form(form_id, request, db, current_user)

@router.delete("/dynamic/form/{form_id}")
def delete_keystone_form(
    form_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return delete_form(form_id, db, current_user)